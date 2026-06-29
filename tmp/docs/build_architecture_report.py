from __future__ import annotations

from datetime import date
from pathlib import Path
import math
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\cibertec\titulacion\new_proyecto\proyecto-radioshack")
OUT_DOC = ROOT / "output" / "documents"
OUT_PDF = ROOT / "output" / "pdf"
TMP = ROOT / "tmp" / "docs"
ASSET_DIR = TMP / "assets"
DOCX_PATH = OUT_DOC / "Arquitectura_Datos_RadioShack_Sustentacion.docx"

for path in (OUT_DOC, OUT_PDF, TMP, ASSET_DIR):
    path.mkdir(parents=True, exist_ok=True)


# standard_business_brief preset, with named report-cover override.
NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "168A8A"
GOLD = "B8860B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D7DEE7"
TEXT = "1F2933"
MUTED = "667085"
WHITE = "FFFFFF"
GREEN = "287D3C"
AMBER = "9A6700"
RED = "A33A3A"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != PAGE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {PAGE_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color="CBD5E1", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_run_font(run, size=11, color=TEXT, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(color)
    run_props.append(underline)
    run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_numbering(doc: Document, fmt: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if fmt == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if fmt == "bullet" else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    # OOXML requires abstract definitions before concrete numbering mappings.
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def add_bullet(doc, text: str, bullet_id: int):
    p = doc.add_paragraph(style="Normal")
    apply_num(p, bullet_id)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_numbered(doc, text: str, number_id: int):
    p = doc.add_paragraph(style="Normal")
    apply_num(p, number_id)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_body(doc, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_caption(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=9, color=MUTED, italic=True)
    return p


def add_callout(doc, title: str, text: str, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    set_table_borders(table, color=accent, size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.5, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, font_size=9, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(header))
        set_run_font(r, size=font_size, color=NAVY, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            color = TEXT
            if str(value).startswith("ALTO"):
                color = RED
            elif str(value).startswith("MEDIO"):
                color = AMBER
            elif str(value).startswith("BAJO"):
                color = GREEN
            set_run_font(r, size=font_size, color=color)
        for idx, width in enumerate(widths):
            tc_pr = cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
    return table


def font(size: int, bold=False):
    path = Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf")
    return ImageFont.truetype(str(path), size)


def wrap_draw(draw, text, box, font_obj, fill=TEXT, align="center", spacing=8):
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 28
    words = text.split()
    lines, current = [], ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font_obj)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    heights = [draw.textbbox((0, 0), line, font=font_obj)[3] for line in lines]
    total_h = sum(heights) + spacing * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=font_obj)[2]
        x = x1 + (x2 - x1 - w) / 2 if align == "center" else x1 + 14
        draw.text((x, y), line, font=font_obj, fill=f"#{fill}")
        y += h + spacing


def rounded_box(draw, box, fill, outline, radius=18, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=width)


def arrow(draw, start, end, color=BLUE, width=6):
    draw.line([start, end], fill=f"#{color}", width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 18
    p1 = (
        end[0] - size * math.cos(angle - math.pi / 6),
        end[1] - size * math.sin(angle - math.pi / 6),
    )
    p2 = (
        end[0] - size * math.cos(angle + math.pi / 6),
        end[1] - size * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, p1, p2], fill=f"#{color}")


def build_logical_architecture(path: Path):
    img = Image.new("RGB", (1800, 900), f"#{WHITE}")
    d = ImageDraw.Draw(img)
    title_font = font(42, True)
    label_font = font(28, True)
    body_font = font(23)
    d.text((70, 35), "Arquitectura logica y flujo de datos", font=title_font, fill=f"#{NAVY}")

    boxes = {
        "source": (70, 210, 360, 470),
        "etl": (440, 150, 760, 530),
        "dw": (840, 150, 1170, 530),
        "api": (1250, 210, 1530, 470),
        "users": (1600, 90, 1760, 310),
        "bi": (1600, 380, 1760, 600),
    }
    rounded_box(d, boxes["source"], "F8FAFC", BLUE)
    rounded_box(d, boxes["etl"], LIGHT_BLUE, BLUE)
    rounded_box(d, boxes["dw"], "E7F6F2", TEAL)
    rounded_box(d, boxes["api"], "FFF7E6", GOLD)
    rounded_box(d, boxes["users"], "F5F3FF", DARK_BLUE)
    rounded_box(d, boxes["bi"], "F5F3FF", DARK_BLUE)

    wrap_draw(d, "RadioShack / ICG SQL Server", boxes["source"], label_font, NAVY)
    d.text((105, 405), "Solo lectura", font=body_font, fill=f"#{MUTED}")

    d.text((525, 185), "ETL Python", font=label_font, fill=f"#{NAVY}")
    for i, line in enumerate(["Extraccion por fecha", "Transformacion pandas", "Controles de calidad", "Carga idempotente"]):
        d.text((480, 260 + i * 58), line, font=body_font, fill=f"#{TEXT}")

    d.text((920, 185), "PostgreSQL", font=label_font, fill=f"#{TEAL}")
    for i, line in enumerate(["stg_ventas", "dim_producto / tienda", "dim_canal / cliente", "fact_ventas"]):
        d.text((875, 260 + i * 58), line, font=body_font, fill=f"#{TEXT}")

    wrap_draw(d, "FastAPI Analytics + Seguridad", boxes["api"], label_font, NAVY)
    wrap_draw(d, "React", boxes["users"], label_font, NAVY)
    wrap_draw(d, "Power BI / Excel", boxes["bi"], label_font, NAVY)

    arrow(d, (360, 340), (440, 340))
    arrow(d, (760, 340), (840, 340))
    arrow(d, (1170, 340), (1250, 340))
    arrow(d, (1530, 300), (1600, 220))
    arrow(d, (1530, 390), (1600, 480))

    d.text((375, 300), "SQL", font=body_font, fill=f"#{MUTED}")
    d.text((770, 300), "UPSERT", font=body_font, fill=f"#{MUTED}")
    d.text((1185, 300), "REST", font=body_font, fill=f"#{MUTED}")

    d.rounded_rectangle((70, 690, 1760, 810), radius=18, fill="#F8FAFC", outline=f"#{MID_GRAY}", width=2)
    d.text((100, 720), "Principio central:", font=label_font, fill=f"#{NAVY}")
    d.text(
        (350, 723),
        "separar el sistema transaccional de la carga analitica y conservar trazabilidad.",
        font=body_font,
        fill=f"#{TEXT}",
    )
    img.save(path, quality=95)


def build_deployment_architecture(path: Path):
    img = Image.new("RGB", (1800, 1020), f"#{WHITE}")
    d = ImageDraw.Draw(img)
    title_font = font(42, True)
    label_font = font(27, True)
    body_font = font(22)
    small = font(19)
    d.text((70, 35), "Arquitectura objetivo de despliegue", font=title_font, fill=f"#{NAVY}")

    rounded_box(d, (70, 180, 300, 420), "F8FAFC", BLUE)
    wrap_draw(d, "Usuarios web", (70, 180, 300, 300), label_font, NAVY)
    wrap_draw(d, "Power BI / Excel", (70, 300, 300, 420), body_font, TEXT)

    rounded_box(d, (420, 155, 770, 445), "FFF7E6", GOLD)
    d.text((495, 190), "Nginx edge", font=label_font, fill=f"#{NAVY}")
    for i, line in enumerate(["HTTPS / TLS", "Proxy /api", "Archivos React", "Cabeceras y limites"]):
        d.text((470, 255 + i * 48), line, font=body_font, fill=f"#{TEXT}")

    d.rounded_rectangle((890, 100, 1710, 760), radius=24, fill="#F8FAFC", outline=f"#{MID_GRAY}", width=4)
    d.text((920, 125), "Servidor Linux - Docker Compose", font=label_font, fill=f"#{NAVY}")

    rounded_box(d, (950, 220, 1270, 440), LIGHT_BLUE, BLUE)
    wrap_draw(d, "Backend FastAPI", (950, 220, 1270, 315), label_font, NAVY)
    wrap_draw(d, "JWT / RBAC / API keys / ETL", (970, 315, 1250, 420), body_font, TEXT)

    rounded_box(d, (1360, 220, 1650, 440), "E7F6F2", TEAL)
    wrap_draw(d, "PostgreSQL", (1360, 220, 1650, 315), label_font, TEAL)
    wrap_draw(d, "Volumen persistente", (1380, 315, 1630, 420), body_font, TEXT)

    rounded_box(d, (950, 520, 1270, 690), "F5F3FF", DARK_BLUE)
    wrap_draw(d, "Logs y metricas", (950, 520, 1270, 610), label_font, NAVY)
    wrap_draw(d, "Alertas / salud", (970, 610, 1250, 680), body_font, TEXT)

    rounded_box(d, (1360, 520, 1650, 690), "FFF1F1", RED)
    wrap_draw(d, "Backup externo", (1360, 520, 1650, 610), label_font, RED)
    wrap_draw(d, "Restore probado", (1380, 610, 1630, 680), body_font, TEXT)

    arrow(d, (300, 300), (420, 300))
    arrow(d, (770, 300), (950, 300))
    arrow(d, (1270, 330), (1360, 330))
    arrow(d, (1505, 440), (1505, 520))
    arrow(d, (1110, 440), (1110, 520))

    rounded_box(d, (70, 580, 430, 790), "F8FAFC", BLUE)
    wrap_draw(d, "Docker Hub privado", (70, 580, 430, 680), label_font, NAVY)
    wrap_draw(d, "Imagenes versionadas por tag y digest", (90, 675, 410, 775), body_font, TEXT)
    arrow(d, (430, 670), (890, 670), TEAL)

    rounded_box(d, (540, 830, 930, 970), "FFF7E6", GOLD)
    wrap_draw(d, "RadioShack / ICG", (540, 830, 930, 900), label_font, NAVY)
    wrap_draw(d, "VPN o IP permitida - solo lectura", (560, 895, 910, 955), small, TEXT)
    arrow(d, (1110, 760), (930, 890), GOLD)

    d.text((70, 850), "Administracion:", font=label_font, fill=f"#{NAVY}")
    d.text((70, 900), "SSH + Docker Compose", font=body_font, fill=f"#{TEXT}")
    d.text((70, 938), "sin escritorio grafico", font=small, fill=f"#{MUTED}")
    img.save(path, quality=95)


def build_scale_ladder(path: Path):
    img = Image.new("RGB", (1800, 700), f"#{WHITE}")
    d = ImageDraw.Draw(img)
    title_font = font(42, True)
    label_font = font(25, True)
    body_font = font(20)
    d.text((70, 35), "Ruta de escalamiento basada en evidencia", font=title_font, fill=f"#{NAVY}")

    stages = [
        ("1. Piloto", "1 host\n2 vCPU / 4 GB\nCompose + backup", BLUE),
        ("2. Departamental", "4 vCPU / 8 GB\nDB separada\ncache y worker", TEAL),
        ("3. Alta demanda", "replicas API\nbalanceador\nread replica", GOLD),
        ("4. Plataforma", "orquestacion\nobservabilidad\nHA multi-zona", DARK_BLUE),
    ]
    x_positions = [80, 505, 930, 1355]
    for idx, (title, body, color) in enumerate(stages):
        x = x_positions[idx]
        box = (x, 190, x + 340, 530)
        rounded_box(d, box, "F8FAFC", color, radius=20, width=4)
        wrap_draw(d, title, (x + 15, 210, x + 325, 300), label_font, color)
        wrap_draw(d, body.replace("\n", " | "), (x + 25, 305, x + 315, 500), body_font, TEXT)
        if idx < len(stages) - 1:
            arrow(d, (x + 340, 360), (x_positions[idx + 1], 360), color)

    d.rounded_rectangle((80, 590, 1695, 660), radius=14, fill="#FFF7E6", outline=f"#{GOLD}", width=2)
    d.text(
        (115, 612),
        "Regla: escalar cuando p95, memoria ETL, conexiones o disponibilidad superen el objetivo; no por intuicion.",
        font=body_font,
        fill=f"#{TEXT}",
    )
    img.save(path, quality=95)


def build_cost_chart(path: Path):
    img = Image.new("RGB", (1800, 850), f"#{WHITE}")
    d = ImageDraw.Draw(img)
    title_font = font(42, True)
    label_font = font(24, True)
    body_font = font(20)
    d.text((70, 35), "Costo mensual referencial de infraestructura", font=title_font, fill=f"#{NAVY}")
    scenarios = [
        ("Base actual\nsolo BD", 45.0, BLUE),
        ("VPS 4 GB\n+ backup", 30.3, TEAL),
        ("VPS 4 GB\n+ Docker Pro", 41.3, GOLD),
        ("VPS 8 GB\n+ backup", 59.1, DARK_BLUE),
        ("BD actual\n+ app VPS", 60.9, RED),
    ]
    chart_left, chart_top, chart_bottom = 130, 150, 670
    max_value = 70
    d.line((chart_left, chart_top, chart_left, chart_bottom), fill=f"#{MID_GRAY}", width=3)
    d.line((chart_left, chart_bottom, 1700, chart_bottom), fill=f"#{MID_GRAY}", width=3)
    for value in range(0, 71, 10):
        y = chart_bottom - (value / max_value) * (chart_bottom - chart_top)
        d.line((chart_left, y, 1700, y), fill="#E9EDF2", width=1)
        d.text((70, y - 12), f"${value}", font=body_font, fill=f"#{MUTED}")
    bar_w = 210
    gap = 90
    for i, (label, value, color) in enumerate(scenarios):
        x = 220 + i * (bar_w + gap)
        h = (value / max_value) * (chart_bottom - chart_top)
        y = chart_bottom - h
        d.rounded_rectangle((x, y, x + bar_w, chart_bottom), radius=12, fill=f"#{color}")
        val = f"${value:.1f}"
        tw = d.textbbox((0, 0), val, font=label_font)[2]
        d.text((x + (bar_w - tw) / 2, y - 42), val, font=label_font, fill=f"#{color}")
        for line_no, line in enumerate(label.split("\n")):
            tw = d.textbbox((0, 0), line, font=body_font)[2]
            d.text((x + (bar_w - tw) / 2, chart_bottom + 20 + line_no * 28), line, font=body_font, fill=f"#{TEXT}")
    d.text(
        (130, 790),
        "Incluye supuestos explicados en el documento; no incluye impuestos, soporte ni tiempo de administracion.",
        font=body_font,
        fill=f"#{MUTED}",
    )
    img.save(path, quality=95)


logical_path = ASSET_DIR / "arquitectura_logica.png"
deployment_path = ASSET_DIR / "arquitectura_despliegue.png"
scale_path = ASSET_DIR / "ruta_escalamiento.png"
cost_path = ASSET_DIR / "costos_referenciales.png"
build_logical_architecture(logical_path)
build_deployment_architecture(deployment_path)
build_scale_ladder(scale_path)
build_cost_chart(cost_path)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = rgb(TEXT)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

caption_style = styles["Caption"]
caption_style.font.name = "Calibri"
caption_style.font.size = Pt(9)
caption_style.font.italic = True
caption_style.font.color.rgb = rgb(MUTED)

bullet_id = create_numbering(doc, "bullet")

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hrun = hp.add_run("ARQUITECTURA DE DATOS | PROYECTO RADIOSHACK")
set_run_font(hrun, size=8.5, color=MUTED, bold=True)

footer = section.footer
fp = footer.paragraphs[0]
add_page_number(fp)


# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(78)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DOCUMENTO DE SUSTENTACION TECNICA")
set_run_font(r, size=11, color=GOLD, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(10)
r = p.add_run("Arquitectura de Datos y\nPlataforma Analitica RadioShack")
set_run_font(r, size=29, color=NAVY, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(34)
r = p.add_run("ETL, modelo dimensional, API segura, dashboard y despliegue contenerizado")
set_run_font(r, size=14, color=DARK_BLUE)

table = doc.add_table(rows=4, cols=2)
set_table_geometry(table, [2700, 6660])
set_table_borders(table, color=MID_GRAY)
cover_rows = [
    ("Tipo de documento", "Informe de arquitectura para proyecto de titulacion"),
    ("Alcance", "Negocio, datos, aplicaciones, tecnologia, seguridad y costos"),
    ("Estado evaluado", "Prototipo funcional local, previo a produccion"),
    ("Fecha", "28 de junio de 2026"),
]
for row, values in zip(table.rows, cover_rows):
    for i, value in enumerate(values):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(value)
        set_run_font(rr, size=10, color=NAVY if i == 0 else TEXT, bold=i == 0)
        if i == 0:
            set_cell_shading(cell, LIGHT_BLUE)

doc.add_paragraph()
add_callout(
    doc,
    "Tesis que sostiene la arquitectura",
    "Separar la fuente transaccional de la carga analitica permite consultar indicadores con seguridad, trazabilidad y menor impacto sobre RadioShack, manteniendo una ruta de crecimiento proporcional al uso real.",
    fill="F8FAFC",
    accent=TEAL,
)

doc.add_page_break()


doc.add_heading("Resumen ejecutivo", level=1)
add_body(
    doc,
    "El proyecto implementa una arquitectura de datos para extraer ventas desde RadioShack/ICG en modo solo lectura, transformar y validar la informacion, cargarla en un almacen PostgreSQL y exponer indicadores mediante una API FastAPI consumida por React, Power BI o Excel. La solucion desacopla el uso analitico del sistema transaccional y evita ejecutar consultas de dashboard directamente sobre la base operativa."
)
add_body(
    doc,
    "El prototipo tiene una base tecnica defendible: ETL incremental por fecha, reproceso idempotente, staging, dimensiones y tabla de hechos, controles fuente-staging-fact, autenticacion JWT, RBAC, API keys con hash y scopes, migraciones, pruebas unitarias y contenedores. No obstante, aun no debe declararse listo para produccion ni certificado bajo ISO/IEC 27001. Faltan controles operativos como TLS, backups probados, CI/CD, monitoreo, rate limiting, gestion de refresh tokens, cierre del registro publico y pruebas de seguridad."
)
add_callout(
    doc,
    "Recomendacion ejecutiva",
    "Para la sustentacion, presentar el despliegue como arquitectura objetivo y demostrar localmente el flujo completo. Un piloto remoto en un unico servidor Linux de 4 GB es razonable si el volumen diario cabe en memoria; 8 GB ofrece mayor margen para pandas y PostgreSQL. Kubernetes no se justifica en esta etapa.",
)

doc.add_heading("Contenido", level=2)
contents_number_id = create_numbering(doc, "decimal")
for item in [
    "Problema, objetivos y valor para el negocio",
    "Arquitectura de datos y aplicaciones",
    "Decisiones tecnologicas y alternativas",
    "Seguridad, normas y brechas",
    "Arquitectura de despliegue",
    "Capacidad, costos y rentabilidad",
    "Escalabilidad y roadmap",
    "Plan de sustentacion y preguntas del jurado",
]:
    add_numbered(doc, item, contents_number_id)


doc.add_heading("1. Problema, objetivos y valor para el negocio", level=1)
doc.add_heading("1.1 Situacion problematica", level=2)
add_body(
    doc,
    "La informacion de ventas reside en una fuente transaccional que solo puede consultarse. Utilizarla directamente para reportes incrementa el acoplamiento, repite calculos, afecta el rendimiento operacional y dificulta conservar reglas de negocio consistentes. Ademas, los campos de fecha-hora no siempre tienen una hora confiable, por lo que el proceso debe trabajar por fecha de negocio y reprocesar una ventana controlada."
)
for text in [
    "Consultas analiticas y operacionales compiten por recursos si se ejecutan en la misma fuente.",
    "Los indicadores pueden diferir si cada herramienta aplica filtros y transformaciones propias.",
    "No existe una capa unica para calidad, trazabilidad, seguridad y consumo externo.",
    "Una base cloud basica puede volverse costosa si recibe consultas repetidas y no agregadas.",
]:
    add_bullet(doc, text, bullet_id)

doc.add_heading("1.2 Objetivo general", level=2)
add_body(
    doc,
    "Disenar e implementar una arquitectura de datos contenerizada que integre ventas de RadioShack en un almacen analitico, aplique controles de calidad y ofrezca indicadores seguros y reutilizables para una aplicacion web, Power BI y Excel."
)

doc.add_heading("1.3 Objetivos especificos", level=2)
objectives_number_id = create_numbering(doc, "decimal")
for text in [
    "Extraer datos sin modificar la fuente y con filtros por fecha.",
    "Normalizar canales, productos, tiendas y clientes mediante un ETL reproducible.",
    "Validar conteos, cantidades y montos entre fuente, staging y hechos.",
    "Modelar ventas en un esquema dimensional optimizado para lectura.",
    "Exponer KPIs, tendencias, ABC, RFM y rankings con filtros y limites.",
    "Proteger usuarios e integraciones con mecanismos de autenticacion distintos.",
    "Definir una ruta de despliegue, seguridad, costo y escalamiento medible.",
]:
    add_numbered(doc, text, objectives_number_id)

doc.add_heading("1.4 Capacidades de negocio habilitadas", level=2)
add_table(
    doc,
    ["Capacidad", "Resultado esperado", "Indicadores"],
    [
        ("Integracion de ventas", "Datos diarios consistentes sin alterar RadioShack", "Filas, monto y cantidad conciliados"),
        ("Analisis comercial", "Vision por fecha, canal, tienda y producto", "Ventas, ticket, unidades, documentos"),
        ("Gestion de portafolio", "Priorizacion de productos por periodo", "ABC, top y bajo movimiento"),
        ("Analisis de clientes", "Segmentacion de comportamiento", "RFM, frecuencia y valor"),
        ("Consumo multicanal", "Una API para web, Power BI y Excel", "JWT y API keys con scope"),
    ],
    [2100, 4100, 3160],
)
add_caption(doc, "Tabla 1. Capacidades de negocio y resultados medibles.")


doc.add_heading("2. Arquitectura de datos y aplicaciones", level=1)
doc.add_picture(str(logical_path), width=Inches(6.35))
add_caption(doc, "Figura 1. Arquitectura logica implementada y flujo de datos.")

doc.add_heading("2.1 Estilo arquitectonico", level=2)
add_body(
    doc,
    "La solucion combina un pipeline ETL con un almacen dimensional y un monolito modular orientado a API. El backend separa routers, servicios, repositorios, schemas y modelos. Esta eleccion mantiene una complejidad adecuada para un proyecto de titulacion y permite evolucionar modulos sin introducir microservicios prematuros."
)
add_table(
    doc,
    ["Capa", "Responsabilidad", "Evidencia actual"],
    [
        ("Fuente", "SQL Server/ICG consultado en modo solo lectura", "Repositorio ICG con consultas parametrizadas"),
        ("ETL", "Extraccion, transformacion, calidad y carga", "Servicios por ventas, productos y tiendas"),
        ("Staging", "Persistencia temporal y diagnostico", "stg_ventas por fecha"),
        ("Warehouse", "Dimensiones y hechos para consulta", "fact_ventas y dimensiones conformadas"),
        ("API", "Seguridad y contratos analiticos", "FastAPI, Pydantic, JWT, RBAC, API keys"),
        ("Presentacion", "Dashboard y consumo externo", "React, Power BI y Excel"),
    ],
    [1500, 3900, 3960],
)
add_caption(doc, "Tabla 2. Capas y responsabilidades.")

doc.add_heading("2.2 Modelo dimensional", level=2)
add_body(
    doc,
    "La tabla fact_ventas representa lineas de venta agregadas por documento, producto, precio y fecha. Se relaciona con dim_producto, dim_tienda, dim_canal y dim_cliente. Este esquema en estrella reduce joins complejos para los consumidores y centraliza codigos de negocio."
)
for text in [
    "Grano declarado: linea analitica de venta por fecha, documento, producto y precio.",
    "Dimensiones conformadas: producto, tienda, canal y cliente.",
    "Medidas: cantidad, precio, descuento, total e IVA.",
    "Indices compuestos: fecha con producto, cliente, tienda y canal.",
    "Pendiente recomendado: migrar montos de float a Numeric/Decimal.",
]:
    add_bullet(doc, text, bullet_id)

doc.add_heading("2.3 Calidad e idempotencia", level=2)
add_body(
    doc,
    "Cada fecha se elimina y reconstruye en staging y hechos. Esta estrategia es idempotente: repetir una fecha produce el mismo estado logico. El proceso compara montos y cantidades entre fuente, staging y fact antes de confirmar. La carga delta reprocesa hoy y ayer para absorber correcciones de una fuente cuya hora no es confiable."
)
add_callout(
    doc,
    "Justificacion para el jurado",
    "El delta no depende de una marca horaria poco confiable. Reprocesar una ventana corta por fecha prioriza consistencia sobre una falsa precision temporal y mantiene el costo acotado.",
    fill="FFF7E6",
    accent=GOLD,
)

doc.add_heading("2.4 Analitica de productos y clientes", level=2)
add_body(
    doc,
    "El ABC debe calcularse dentro de un horizonte de analisis. Para decisiones operativas se recomienda 90 o 180 dias; para planeamiento, doce meses moviles. El historico completo es util como comparacion, pero sesga la clasificacion hacia productos antiguos o descatalogados. El RFM tambien depende del periodo y de los filtros de canal o tienda."
)


doc.add_heading("3. Decisiones tecnologicas y alternativas", level=1)
add_body(
    doc,
    "Las decisiones se evaluaron por adecuacion al problema, costo, mantenibilidad, experiencia disponible y ruta de crecimiento. Elegir tecnologia no significa que las alternativas sean incorrectas, sino que ofrecen menor ajuste al contexto actual."
)
add_table(
    doc,
    ["Decision", "Por que se elige", "Alternativa y trade-off"],
    [
        ("Python + pandas", "Ecosistema ETL, tratamiento tabular, rapidez de desarrollo y librerias abiertas [5].", "Java/Spark: mayor robustez distribuida, pero complejidad excesiva para el volumen actual."),
        ("FastAPI", "API-first, validacion Pydantic, OpenAPI y Swagger automaticos [7].", "Django: mas completo para aplicaciones monoliticas, pero mayor superficie para una API analitica."),
        ("PostgreSQL", "ACID, integridad, ventanas, indices, materializacion y licencia abierta [6].", "SQL Server: continuidad con la fuente, pero mayor acoplamiento y posible costo de licencia."),
        ("React + TypeScript", "Componentes reutilizables y tipado para dashboards [8].", "Angular: estructura mas prescriptiva, con mayor curva y peso para este alcance."),
        ("Vite", "Ciclo de desarrollo rapido y build optimizado de activos estaticos [9].", "Webpack manual: flexible, pero requiere mas configuracion y mantenimiento."),
        ("Docker", "Empaquetado reproducible, aislamiento y portabilidad [10].", "Instalacion nativa: menos capas, pero mayor deriva entre ambientes."),
        ("Docker Compose", "Define servicios, redes y volumenes en un YAML; adecuado para un host [11].", "Kubernetes: alta disponibilidad y orquestacion, no justificadas por la escala actual."),
        ("Nginx", "Punto unico para HTTPS, proxy, archivos estaticos y limites [12].", "Exponer Uvicorn: mas simple, pero sin capa edge ni terminacion TLS."),
    ],
    [1500, 3980, 3880],
    font_size=8.4,
)
add_caption(doc, "Tabla 3. Registro resumido de decisiones arquitectonicas.")

doc.add_heading("3.1 Por que un servidor Linux y no un escritorio remoto", level=2)
add_body(
    doc,
    "El termino correcto es servidor Linux o VPS. La administracion se realiza por SSH; un escritorio grafico consume memoria, amplia la superficie de ataque y no aporta valor al runtime. Docker Engine y Compose funcionan como servicios del sistema. El despliegue consiste en autenticar el registro, descargar imagenes versionadas, aplicar migraciones y levantar el stack."
)

doc.add_heading("3.2 Docker Hub como registro", level=2)
add_body(
    doc,
    "Docker Hub permite distribuir imagenes inmutables en lugar de copiar codigo al servidor. El plan Personal incluye un repositorio privado y tiene limites de pulls; dos imagenes privadas separadas pueden requerir Docker Pro o un registro alternativo. En junio de 2026 Docker Pro figura en USD 11/mes [14]. Las imagenes deben publicarse sin secretos y referenciarse por tag versionado o digest."
)


doc.add_heading("4. Seguridad, normas y brechas", level=1)
add_callout(
    doc,
    "Declaracion correcta",
    "El proyecto esta alineado parcialmente con buenas practicas de ISO/IEC 27001:2022, NIST CSF 2.0, OWASP ASVS 5.0.0 y CIS Docker Benchmark. No esta certificado ni puede afirmarse cumplimiento total sin SGSI, evaluacion formal, evidencia y auditoria.",
    fill="FFF7E6",
    accent=GOLD,
)
add_body(
    doc,
    "ISO/IEC 27001 define requisitos para un sistema de gestion de seguridad de la informacion que abarca personas, procesos y tecnologia [1]. NIST CSF 2.0 organiza la gestion del riesgo en Govern, Identify, Protect, Detect, Respond y Recover [2]. OWASP ASVS proporciona requisitos verificables para controles tecnicos de aplicaciones web [3]. CIS publica una guia de configuracion segura para Docker [4]."
)

doc.add_heading("4.1 Controles implementados", level=2)
add_table(
    doc,
    ["Dominio", "Control actual", "Evidencia", "Estado"],
    [
        ("Acceso", "JWT, roles, permisos y API keys con scope", "security.*, PermissionChecker, api_clients", "PARCIAL"),
        ("Credenciales", "Passwords hasheadas; API keys solo como hash", "pwdlib y SHA-256 de claves aleatorias", "IMPLEMENTADO"),
        ("Datos fuente", "Usuario de consulta sin modificacion", "Conexion ICG separada", "IMPLEMENTADO"),
        ("Integridad", "Conciliacion fuente-staging-fact", "Controles de monto y cantidad", "IMPLEMENTADO"),
        ("Auditoria", "Trace ID y redaccion de secretos", "AuditMiddleware y sanitizer", "PARCIAL"),
        ("Segregacion", "ETL, seguridad y analytics con permisos", "coolbox.etl.execute / analytics.read", "IMPLEMENTADO"),
        ("Cadena de suministro", "Versiones Python fijadas y package-lock", "requirements.txt / package-lock.json", "PARCIAL"),
        ("Despliegue", "Red y volumenes Compose", "compose de produccion", "OBJETIVO"),
    ],
    [1550, 2850, 3300, 1660],
    font_size=8.2,
)
add_caption(doc, "Tabla 4. Mapeo de controles actuales. Estado cualitativo, no auditoria.")

doc.add_heading("4.2 Brechas antes de produccion", level=2)
add_table(
    doc,
    ["Prioridad", "Brecha", "Tratamiento propuesto", "Evidencia de cierre"],
    [
        ("ALTO", "Registro de usuarios actualmente publico", "Restringir a Admin o flujo de invitacion", "Prueba 403 sin permiso"),
        ("ALTO", "JWT en localStorage y sin refresh token", "Cookie HttpOnly/Secure, access corto y rotacion", "Pruebas login-refresh-logout"),
        ("ALTO", "Sin HTTPS en arquitectura actual", "Nginx + certificado + redireccion HTTP", "Escaneo TLS y evidencia navegador"),
        ("ALTO", "Sin backup/restore probado", "Backup cifrado externo y simulacro mensual", "Acta de restauracion y RPO/RTO"),
        ("ALTO", "Sin rate limiting", "Limites por IP, usuario y API key", "Prueba de respuesta 429"),
        ("MEDIO", "Contenedores ejecutan como root", "Usuario no privilegiado y filesystem restringido", "docker inspect / Dockerfile"),
        ("MEDIO", "Sin CI/CD ni escaneo de imagen", "Build, tests, SBOM, scan y push por tag", "Pipeline exitoso y reporte"),
        ("MEDIO", "Cobertura de pruebas limitada", "Integracion ETL, seguridad y contratos", "Reporte de cobertura"),
        ("MEDIO", "Sin monitoreo ni alertas", "Metricas, logs centralizados y healthchecks", "Dashboard y alerta de prueba"),
    ],
    [1100, 2640, 3700, 1920],
    font_size=8,
)
add_caption(doc, "Tabla 5. Backlog minimo de endurecimiento.")

doc.add_heading("4.3 Objetivo de seguridad defendible", level=2)
for text in [
    "Objetivo inicial: OWASP ASVS nivel 1 completo y controles seleccionados de nivel 2 para autenticacion.",
    "Aplicar principio de minimo privilegio en fuente, warehouse, API keys y contenedores.",
    "Definir inventario de activos, riesgos, responsables y tratamiento, alineado a NIST Govern/Identify.",
    "Documentar respuesta a incidentes, backups, RPO y RTO, alineado a Detect/Respond/Recover.",
    "No almacenar secretos en Docker Hub, imagenes, repositorios ni documentos.",
]:
    add_bullet(doc, text, bullet_id)


doc.add_heading("5. Arquitectura de despliegue", level=1)
doc.add_picture(str(deployment_path), width=Inches(6.35))
add_caption(doc, "Figura 2. Despliegue objetivo en un servidor Linux con Docker Compose.")
add_body(
    doc,
    "Solo Nginx publica 80/443. FastAPI y PostgreSQL permanecen en redes internas. La base no expone 5432 a Internet. La conexion a RadioShack requiere VPN, red corporativa o IP publica autorizada y credenciales de solo lectura. Los backups deben salir del mismo host para evitar que una falla del servidor destruya datos y copias."
)

doc.add_heading("5.1 Servicios propuestos en Compose", level=2)
add_table(
    doc,
    ["Servicio", "Funcion", "Exposicion", "Persistencia"],
    [
        ("nginx", "TLS, frontend, proxy /api y limites", "80/443", "Certificados/config"),
        ("backend", "API, analytics y ETL", "Solo red interna", "Sin estado"),
        ("postgres", "Warehouse y seguridad", "Solo red interna", "Volumen + backup"),
        ("backup", "Dump y copia externa programada", "Sin puerto", "Destino externo"),
        ("monitoring (opcional)", "Metricas, logs y alertas", "VPN o interno", "Segun herramienta"),
    ],
    [1600, 3500, 2100, 2160],
)
add_caption(doc, "Tabla 6. Servicios y fronteras de exposicion.")

doc.add_heading("5.2 Flujo de despliegue", level=2)
deployment_number_id = create_numbering(doc, "decimal")
for text in [
    "Ejecutar tests y construir imagenes con tags inmutables, por ejemplo 1.0.0.",
    "Escanear dependencias e imagenes; generar SBOM.",
    "Publicar backend y frontend/edge en registro privado.",
    "Conectar por SSH al servidor y descargar las imagenes.",
    "Ejecutar alembic upgrade head como tarea controlada.",
    "Levantar Compose y esperar healthchecks.",
    "Ejecutar smoke tests de login, analytics y conectividad ICG.",
    "Conservar la version anterior para rollback.",
]:
    add_numbered(doc, text, deployment_number_id)


doc.add_heading("6. Capacidad, costos y rentabilidad", level=1)
doc.add_heading("6.1 Estimacion inicial de recursos", level=2)
add_table(
    doc,
    ["Componente", "RAM normal estimada", "Pico / observacion"],
    [
        ("Nginx edge + frontend", "40-120 MB", "Depende de conexiones y cache"),
        ("FastAPI 1 worker", "250-500 MB", "Incluye pandas, numpy y drivers"),
        ("PostgreSQL", "600 MB-2 GB", "Cache, work_mem y consultas"),
        ("ETL diario", "0.5-3 GB adicionales", "Depende de filas y columnas del dia"),
        ("Sistema Linux", "300-700 MB", "Docker, kernel y agentes"),
    ],
    [2300, 2300, 4760],
)
add_caption(doc, "Tabla 7. Estimaciones orientativas; deben reemplazarse por medicion.")
add_body(
    doc,
    "Un host de 2 vCPU y 4 GB puede servir como piloto si el ETL se ejecuta sin alta concurrencia y el volumen diario es moderado. Para reducir riesgo de out-of-memory durante pandas, 4 vCPU y 8 GB es una base mas comoda. El disco debe dimensionarse como minimo para datos activos, indices, temporales y dos copias de respaldo; una regla inicial es tres veces el tamano actual medido."
)

doc.add_heading("6.2 Escenarios de costo", level=2)
doc.add_picture(str(cost_path), width=Inches(6.35))
add_caption(doc, "Figura 3. Comparacion mensual referencial, precios consultados el 28/06/2026.")
add_table(
    doc,
    ["Escenario", "Supuesto mensual", "Total aprox.", "Uso recomendado"],
    [
        ("Actual", "Base cloud reportada por el autor", "USD 45", "Desarrollo; costo alto si solo cubre BD"),
        ("Piloto unico", "VPS 4 GB USD 24 + backup 20% + dominio USD 1.5", "USD 30.3", "Demo remota y pocos usuarios"),
        ("Piloto privado", "Piloto + Docker Pro USD 11", "USD 41.3", "Dos imagenes privadas en Docker Hub"),
        ("Host con margen", "VPS 8 GB USD 48 + backup 20% + dominio", "USD 59.1", "ETL mas grande y mayor estabilidad"),
        ("Separado", "BD actual USD 45 + VPS 2 GB USD 12 + backup/dominio", "USD 60.9", "DB administrada y app separada"),
    ],
    [1600, 3890, 1470, 2400],
    font_size=8.2,
)
add_caption(doc, "Tabla 8. No incluye impuestos, soporte, egress ni trabajo operativo.")
add_body(
    doc,
    "DigitalOcean publica USD 24/mes para 2 vCPU, 4 GB y 80 GB SSD, y USD 48/mes para 4 vCPU, 8 GB y 160 GB SSD; el backup semanal porcentual se publica al 20% del costo del servidor [13]. Son referencias comparables, no una recomendacion contractual."
)

doc.add_heading("6.3 Rentabilidad y punto de equilibrio", level=2)
add_body(
    doc,
    "La plataforma no genera ingresos por si sola; crea ahorro y capacidad de decision. La forma defendible de cuantificarla es medir horas evitadas, errores reducidos y decisiones comerciales mejoradas."
)
add_callout(
    doc,
    "Formula",
    "Beneficio mensual = horas ahorradas x costo-hora cargado + costo de errores evitados + margen incremental atribuible. ROI = (beneficio - costo) / costo.",
    fill="E7F6F2",
    accent=TEAL,
)
# Keep the ROI scenarios together instead of leaving an orphaned table header.
doc.add_page_break()
add_table(
    doc,
    ["Escenario", "Supuesto de ahorro", "Beneficio", "Infra", "Resultado"],
    [
        ("Conservador", "1 analista x 6 h x USD 6", "USD 36", "USD 30", "USD 6 netos; 1.2x"),
        ("Base", "2 analistas x 8 h x USD 8", "USD 128", "USD 41", "USD 87 netos; 3.1x"),
        ("Alto", "3 analistas x 12 h x USD 10", "USD 360", "USD 60", "USD 300 netos; 6.0x"),
    ],
    [1550, 2930, 1500, 1300, 2080],
)
add_caption(doc, "Tabla 9. Ejemplos para la metodologia; reemplazar por datos reales.")
add_body(
    doc,
    "Con un costo de USD 41/mes y un costo-hora de USD 8, el punto de equilibrio es 5.1 horas ahorradas por mes. Esta cifra es facil de validar durante un piloto comparando el tiempo anterior y posterior para preparar el mismo reporte."
)


doc.add_heading("7. Escalabilidad y roadmap", level=1)
doc.add_picture(str(scale_path), width=Inches(6.35))
add_caption(doc, "Figura 4. Escalamiento progresivo sin sobredisenar.")

doc.add_heading("7.1 Indicadores que disparan escalamiento", level=2)
add_table(
    doc,
    ["Metrica", "Objetivo inicial", "Accion si se supera"],
    [
        ("Latencia API p95", "< 2 s en dashboard", "EXPLAIN, agregados, cache; luego mas CPU"),
        ("Memoria ETL", "< 70% del host", "Chunks, tipos compactos o mas RAM"),
        ("CPU sostenida", "< 70% fuera del batch", "Optimizar consulta o separar worker"),
        ("Conexiones DB", "< 70% del limite", "Ajustar pool, cache o separar cargas"),
        ("Disponibilidad", ">= 99% piloto", "Healthchecks, alertas, replicas"),
        ("RPO / RTO", "24 h / 4 h piloto", "Backups mas frecuentes y restore automatizado"),
    ],
    [2100, 2650, 4610],
)
add_caption(doc, "Tabla 10. Metricas de capacidad propuestas.")

# Start the roadmap as one readable block instead of repeating its header alone.
doc.add_page_break()
doc.add_heading("7.2 Roadmap recomendado", level=2)
add_table(
    doc,
    ["Fase", "Entregables", "Criterio de salida"],
    [
        ("0. Titulacion", "Arquitectura, demo local, pruebas, costos y riesgos", "Sustentacion reproducible"),
        ("1. Hardening", "HTTPS, registro cerrado, refresh token, rate limit, backup", "Checklist critico cerrado"),
        ("2. Piloto", "Servidor Linux, Compose, dominio, monitoreo basico", "2-4 semanas sin incidentes"),
        ("3. Optimizacion", "EXPLAIN, agregado diario, cache de filtros, chunks", "p95 y RAM dentro de objetivo"),
        ("4. Escala", "DB separada, worker, replicas y observabilidad", "Demanda medible lo justifica"),
    ],
    [1500, 4650, 3210],
)
add_caption(doc, "Tabla 11. Roadmap incremental.")

doc.add_heading("7.3 Mejoras tecnicas priorizadas", level=2)
for text in [
    "Crear agg_ventas_diarias y actualizarla al terminar cada fecha ETL.",
    "Ejecutar ABC y RFM sobre agregados o vistas materializadas cuando el volumen lo requiera.",
    "Procesar extracciones grandes por chunks para acotar memoria.",
    "Particionar fact_ventas por mes cuando la medicion demuestre beneficio.",
    "Separar ETL en un worker programado para no competir con la API.",
    "Agregar cache de filtros y KPIs con TTL corto.",
    "Incorporar pruebas de carga con volumen representativo.",
]:
    add_bullet(doc, text, bullet_id)


doc.add_heading("8. Plan de sustentacion", level=1)
doc.add_heading("8.1 Narrativa de 10 minutos", level=2)
add_table(
    doc,
    ["Minuto", "Mensaje", "Evidencia"],
    [
        ("0-1", "Problema y restriccion de solo lectura", "Contexto de RadioShack"),
        ("1-3", "Flujo ETL, staging y modelo dimensional", "Figura 1 y conciliaciones"),
        ("3-5", "Analytics y decisiones de tecnologia", "Endpoints, ABC/RFM y Tabla 3"),
        ("5-7", "Seguridad y consumo web/BI", "JWT, RBAC, API keys y brechas"),
        ("7-9", "Despliegue, costo y escalabilidad", "Figuras 2-4"),
        ("9-10", "Valor, limites y siguiente fase", "ROI y roadmap"),
    ],
    [1050, 4650, 3660],
)
add_caption(doc, "Tabla 12. Guion sugerido de exposicion.")

doc.add_heading("8.2 Preguntas probables del jurado", level=2)
qa = [
    ("¿Por que copiar los datos?", "Para desacoplar analitica del sistema operacional, aplicar calidad y evitar impacto sobre la fuente."),
    ("¿Por que no consultar SQL Server directamente?", "La fuente es de solo lectura, operacional y no controla las reglas comunes del dashboard."),
    ("¿Por que reprocesar hoy y ayer?", "La hora no es confiable; una ventana por fecha absorbe correcciones con costo predecible."),
    ("¿Por que PostgreSQL?", "Integridad ACID, SQL analitico, indices, materializacion, extensibilidad y ausencia de licencia propietaria."),
    ("¿Por que no microservicios?", "El dominio y equipo no justifican costos distribuidos; el monolito modular conserva limites claros."),
    ("¿Por que Docker Compose y no Kubernetes?", "Hay un host y baja escala; Compose aporta reproducibilidad sin el costo operativo de un cluster."),
    ("¿Cumple ISO 27001?", "No se afirma certificacion. Se demuestra alineacion parcial y un backlog para un SGSI y verificacion formal."),
    ("¿Como se recupera de una falla?", "Objetivo: backup externo, restore probado, imagen anterior y migracion reversible; aun es brecha de produccion."),
    ("¿Como escala?", "Primero agregados, cache y chunks; luego separacion de DB/worker y finalmente replicas si las metricas lo exigen."),
    ("¿Es rentable?", "Se calcula por horas y errores evitados; el piloto mide el antes/despues y valida el punto de equilibrio."),
]
for question, answer in qa:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(question)
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.space_after = Pt(7)
    r2 = p2.add_run(answer)
    set_run_font(r2, size=10.5, color=TEXT)

doc.add_heading("8.3 Demostracion sugerida", level=2)
demo_number_id = create_numbering(doc, "decimal")
for text in [
    "Mostrar el login y el usuario actual.",
    "Ejecutar filtros de fecha, tienda y canal.",
    "Comparar KPI, evolucion y top productos.",
    "Mostrar ABC para 12 meses y luego para una tienda.",
    "Crear una API key y explicar que el secreto solo se muestra una vez.",
    "Consultar un endpoint con X-API-Key desde Power Query o Postman.",
    "Mostrar una ejecucion delta y sus conteos de control.",
    "Cerrar con docker compose config y el diagrama objetivo, sin depender de Internet.",
]:
    add_numbered(doc, text, demo_number_id)


doc.add_heading("9. Criterios de aceptacion para produccion", level=1)
checks = [
    ("Seguridad", "HTTPS, registro restringido, refresh token seguro, rate limit y secretos fuera de imagen"),
    ("Datos", "Migraciones aplicadas, conciliaciones, Decimal y politicas de retencion"),
    ("Continuidad", "Backup externo, restauracion probada, RPO/RTO y rollback"),
    ("Operaciones", "Healthchecks, alertas, logs, espacio en disco y rotacion"),
    ("Calidad", "Tests unitarios, integracion, contratos, carga y seguridad"),
    ("Despliegue", "Imagenes privadas/versionadas, SBOM, scan y CI/CD"),
    ("Red", "Solo 80/443, PostgreSQL privado, VPN/IP permitida a ICG"),
    ("Gobierno", "Responsables, inventario, riesgos, accesos y evidencia"),
]
add_table(doc, ["Area", "Criterio de salida"], checks, [1800, 7560])
add_caption(doc, "Tabla 13. Definition of Done para autorizar produccion.")

add_callout(
    doc,
    "Decision final sugerida",
    "El proyecto esta listo para sustentarse como prototipo funcional y arquitectura objetivo. Antes de exponerlo a Internet, ejecutar la fase de hardening y un piloto controlado. La produccion es una decision posterior basada en uso, costo y riesgo, no un requisito para validar la arquitectura.",
    fill="E7F6F2",
    accent=TEAL,
)


doc.add_heading("Referencias", level=1)
sources = [
    ("[1] ISO/IEC 27001:2022 - Information security management systems.", "https://www.iso.org/standard/27001"),
    ("[2] NIST Cybersecurity Framework 2.0.", "https://www.nist.gov/cyberframework"),
    ("[3] OWASP Application Security Verification Standard 5.0.0.", "https://owasp.org/www-project-application-security-verification-standard/"),
    ("[4] CIS Docker Benchmark.", "https://www.cisecurity.org/benchmark/docker"),
    ("[5] Python Software Foundation - About Python.", "https://www.python.org/about/"),
    ("[6] PostgreSQL - About.", "https://www.postgresql.org/about/"),
    ("[7] FastAPI - Features.", "https://fastapi.tiangolo.com/features/"),
    ("[8] React - Quick Start.", "https://react.dev/learn"),
    ("[9] Vite - Why Vite.", "https://vite.dev/guide/why.html"),
    ("[10] Docker - What is Docker?", "https://docs.docker.com/get-started/docker-overview/"),
    ("[11] Docker Compose documentation.", "https://docs.docker.com/compose/"),
    ("[12] Nginx official site and documentation.", "https://nginx.org/"),
    ("[13] DigitalOcean Droplet Pricing, consulted 2026-06-28.", "https://www.digitalocean.com/pricing/droplets"),
    ("[14] Docker plans and Docker Hub usage limits, consulted 2026-06-28.", "https://www.docker.com/pricing/"),
]
for label, url in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_hyperlink(p, label, url)

doc.add_heading("Anexo A. Evidencias tecnicas actuales", level=1)
add_table(
    doc,
    ["Evidencia", "Estado observado"],
    [
        ("Backend", "FastAPI, SQLAlchemy, Alembic, pandas y PostgreSQL"),
        ("Frontend", "React 19, TypeScript, Vite 8, Zustand y Recharts"),
        ("ETL", "Fecha, rango, delta, staging, controles y carga por lotes"),
        ("Analytics", "Ventas, productos, clientes, ABC, RFM y filtros"),
        ("Seguridad", "JWT, RBAC, API clients, hash, scope y auditoria redactada"),
        ("Pruebas", "14 pruebas unitarias aprobadas al cierre de la revision"),
        ("Migraciones", "Una cabeza Alembic: d91a0d67c8fe"),
        ("Contenedores", "Dockerfiles y Compose validos; falta stack unificado con Nginx"),
    ],
    [2300, 7060],
)
add_caption(doc, "Tabla 14. Evidencia obtenida del codigo local revisado.")

doc.add_heading("Anexo B. Supuestos y limites", level=1)
for text in [
    "Las cifras de RAM son estimaciones y deben validarse con docker stats y una carga representativa.",
    "Los costos no incluyen impuestos, soporte, egress, dominio real ni horas de administracion.",
    "La disponibilidad no puede garantizarse en un unico host.",
    "El acceso remoto a RadioShack depende de conectividad autorizada fuera del codigo.",
    "La rentabilidad debe reemplazar los ejemplos por salarios, tiempos y errores reales.",
    "La alineacion normativa no equivale a certificacion ni auditoria independiente.",
]:
    add_bullet(doc, text, bullet_id)

doc.core_properties.title = "Arquitectura de Datos y Plataforma Analitica RadioShack"
doc.core_properties.subject = "Documento de sustentacion tecnica para proyecto de titulacion"
doc.core_properties.keywords = "arquitectura de datos, ETL, FastAPI, PostgreSQL, Docker, seguridad"
doc.core_properties.author = "Proyecto RadioShack"

doc.save(DOCX_PATH)
print(DOCX_PATH)
